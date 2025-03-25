# from Method_Functions import 绘制点位分布, 绘制污染点位分布V2, 绘制保存异常点位
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from datetime import datetime
from docx.oxml.ns import qn  # 用于设置中文字体
from pathlib import Path
import geopandas as gpd
from docx.shared import RGBColor


def set_heading_style(heading, level=1):
    """
    设置一级标题样式：
    - 字体为黑体
    - 字号为三号
    - 加粗
    - 顶格（左对齐）

    :heading: 段落对象
    :text: 标题文本
    """
    # 设置标题的字体和字号
    run = heading.runs[0]  # 获取标题的第一个 Run 对象
    if level == 1:
        run.font.size = Pt(16)  # 设置字号为三号字（16磅）
        run.bold = True  # 设置加粗
        run.font.name = "Times New Roman"  # 设置字体为黑体
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")  # 设置中文为黑体
        run.font.color.rgb = RGBColor(0, 0, 0)

        # 设置段落格式（顶格、左对齐）
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
    elif level == 2:
        run.font.size = Pt(12)  #
        run.bold = True  # 设置加粗
        run.font.name = "SimSong"  # 设置字体为黑体
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 设置中文为黑体
        run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置段落格式（顶格、左对齐）
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐


def set_paragraph_style(
    paragraph,
    is_chinese=True,
    font_size=12,
    line_spacing=1.5,
    space_before=6,
    space_after=6,
):
    """
    设置段落的字体、字号、行距和段落间距
    :param paragraph: 段落对象
    :param is_chinese: 是否是中文，True 表示中文，False 表示英文
    :param font_size: 字号（磅）
    :param line_spacing: 行距
    :param space_before: 段前间距
    :param space_after: 段后间距
    """
    # 设置首行缩进 2 字符

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Pt(24)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 设置两端对齐
    paragraph_format.line_spacing = line_spacing  # 设置行距
    paragraph_format.space_before = Pt(space_before)  # 段前间距
    paragraph_format.space_after = Pt(space_after)  # 段后间距

    # 设置 Run 的字体属性
    for run in paragraph.runs:
        if is_chinese:
            run.font.name = "Times New Roman"  # 中文设置
            run.font.color.rgb = RGBColor(0, 0, 0)
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 设置黑体
        else:
            run.font.name = "Times New Roman"  # 英文设置
            run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(font_size)  # 设置字号


def add_paragraph(doc, text):
    """
    插入自定义格式的段落：小四字号、宋体、1.5倍行距、两端对齐、首行缩进2字符。

    :param doc: Document 对象
    :param text: 要插入的文本
    """
    # 插入段落
    paragraph = doc.add_paragraph(text)

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.first_line_indent = Pt(
        24
    )  # 首行缩进2字符（每个字符12磅，2字符即24磅）

    # 设置字体样式
    for run in paragraph.runs:
        run.font.name = "Times New Roman"  # 设置字体为宋体
        run.font.size = Pt(12)  # 小四字号（12磅）
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 确保设置为宋体

    return paragraph


def add_table_header(doc, text):
    # 插入段落
    paragraph = doc.add_paragraph(text)

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.space_before = Pt(6)  # 段前间距
    paragraph_format.space_after = Pt(0)  # 段后间距
    # 设置字体样式
    for run in paragraph.runs:
        run.font.name = "Times New Roman"  # 设置字体为宋体
        run.font.size = Pt(12)  # 小四字号（12磅）
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 确保设置为宋体

    return paragraph


def add_pic_header(doc, text):
    # 插入段落
    paragraph = doc.add_paragraph(text)

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.space_before = Pt(0)  # 段前间距
    paragraph_format.space_after = Pt(6)  # 段后间距
    # 设置字体样式
    for run in paragraph.runs:
        run.font.name = "Times New Roman"  # 设置字体为宋体
        run.font.size = Pt(12)  # 小四字号（12磅）
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 确保设置为宋体

    return paragraph


from PIL import Image
from docx import Document
from docx.shared import Inches


def insert_image(doc, image_path, width=None, height=None):
    """
    插入图片，并支持设置宽度和高度，图片默认居中显示。

    :param doc: Document 对象
    :param image_path: 图片路径
    :param width: 图片宽度（可选，单位为英寸，默认不设置）
    :param height: 图片高度（可选，单位为英寸，默认不设置）
    """
    if not Path(image_path).exists():
        print(f"图片路径无效：{image_path}")
        return
    # 创建段落
    paragraph = doc.add_paragraph()

    # 在段落中插入图片
    run = paragraph.add_run()

    # 设置图片宽度和高度
    if width and height:
        run.add_picture(
            image_path, width=Inches(width), height=Inches(height)
        )  # 设置宽高
    elif width:
        run.add_picture(image_path, width=Inches(width))  # 只设置宽度
    elif height:
        run.add_picture(image_path, height=Inches(height))  # 只设置高度
    else:
        run.add_picture(image_path)  # 默认不设置，按原图大小插入

    # 设置图片所在段落居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中显示图片


def apply_heading_style_to_all_titles(doc):
    """
    遍历文档中的所有一级标题，应用标题样式。

    :param doc: Document 对象
    """
    for paragraph in doc.paragraphs:
        # 判断该段落是否为一级标题
        if paragraph.style.name == "Heading 1":  # 默认为一级标题的样式名为 "Heading 1"
            set_heading_style(paragraph, level=1)  # 设置标题样式
        elif paragraph.style.name == "Heading 2":
            set_heading_style(paragraph, level=2)
        # elif paragraph.style.name == "Normal":
        #     set_paragraph_style(paragraph)


def auto_report_EN():
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION
    from datetime import datetime
    from docx.oxml.ns import qn  # 用于设置中文字体
    from pathlib import Path
    import geopandas as gpd

    # 创建文档
    doc = Document()

    # 封面标题

    # 添加空段落，使标题位于页面中部
    for _ in range(8):  # 根据页面长度适当调整数量
        i = doc.add_paragraph()
        i.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = doc.add_paragraph()
    title_run = title.add_run("Microturbulent Pollution Investigation Report")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    title_run.font.name = "Times New Roman"  # 设置中文字体
    title_run.font.size = Pt(40)  # 设置字体大小为三号
    title_run.bold = True  # 加粗
    r = title_run._element  # 处理中文字体
    r.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # 添加单位名称
    # 添加空段落，使标题位于页面中部
    for _ in range(10):  # 根据页面长度适当调整数量
        i = doc.add_paragraph()
        i.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    organization_paragraph = doc.add_paragraph()
    organization_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    organization_run = organization_paragraph.add_run(
        "XX environmental technology company"
    )
    organization_run.font.name = "Times New Roman"
    organization_run.font.size = Pt(12)
    r = organization_run._element  # 处理中文字体
    r.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # 添加时间
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    date_run = date_paragraph.add_run(datetime.now().strftime("%Y-%m"))
    date_run.font.name = "Times New Roman"
    date_run.font.size = Pt(16)
    r = date_run._element  # 处理中文字体
    r.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # 添加分节符
    doc.add_section(WD_SECTION.NEW_PAGE)

    # 添加正文示例
    doc.add_heading("0 Clarification", level=1)
    add_paragraph(
        doc,
        "Bound by a confidentiality agreement, this report only provides examples and methods of arrangement, and does not contain real content.",
    )
    add_paragraph(
        doc,
        "Below we show how to insert multi-level headings, paragraphs, tables, images, etc. through a partial demo. Adjust the code according to different needs to generate a report that meets the requirements.",
    )
    doc.add_heading("1 ", level=1)
    add_paragraph(
        doc,
        """
    一份完整的调查报告通常包括以下内容：
    1. 项目背景：简要介绍调查项目的背景和目的；
    2. 现场调查：对调查现场的基本情况、调查方法和调查结果进行描述；
    3. 数据分析：对调查数据进行评估和分析，确定污染源区和污染羽；
    4. 结果展示：通过图表、表格等形式展示调查结果；
    5. 结论与建议：总结调查结果，提出后续工作建议。
    """,
    )

    add_paragraph(
        doc,
        "",
    )

    doc.add_heading("2 Adding different levels of headings", level=1)
    doc.add_heading("2.1 Secondary headings", level=2)
    doc.add_heading("2.1.1 Tertiary heading", level=3)

    insert_image(doc, "./ref/微扰动原理.png", width=4)
    add_pic_header(doc, "图3-1 微扰动调查技术原理图")

    # 填充表格
    import random

    table_header = [
        "ID",
        "Rn(Bq/m³)",
        "VOCs(ppb)",
        "CO2(ppm)",
        "O2(%)",
        "CH4(%)",
        "FG(copies/g)",
    ]
    add_table_header(doc, "Table4-1 Microperturbation survey site data")
    table = doc.add_table(rows=7, cols=15)
    table.style = "Table Grid"
    for i, col_name in enumerate(table_header):
        table.cell(0, i).text = col_name
        for j in range(1, 15):
            table.cell(0, j).text = str(random.randint(10, 100))

    doc.add_heading("5 数据评估与结果分析", level=1)
    doc.add_heading("5.1 判定原则", level=2)
    add_paragraph(
        doc,
        "调查人员须具备一定专业知识和工作经验，推荐综合所有指标，以K-MEANS聚类法、参考值法等，识别环境背景区域，确定微扰动调查指标的环境环境背景值。与环境环境背景值比较，其中氡气明显异常低，是微扰动调查判定污染源区的必要条件。以如下其他四个辅助条件圈定污染羽，并可辅以判定污染源区：①VOCs明显异常高；②CH4明显异常高；③CO2明显异常高、O2明显异常低；④功能基因明显异常高。",
    )
    add_paragraph(doc, "（1）污染源区判定原则")
    add_paragraph(
        doc,
        "①必要条件满足，且其他四个辅助条件均满足，可直接依据微扰动调查成果，将点位所在区域判定为污染源区；",
    )
    add_paragraph(
        doc,
        "②必要条件满足，且其他辅助条件满足三个，可结合采样检测进行验证，当表层土壤和地下水中污染物浓度含量均超过相关标准限值要求时，将点位所在区域判定为污染源区；",
    )
    add_paragraph(
        doc,
        "③必要条件满足，且其他辅助条件满足两个及以下，须结合采样检测进行验证，尤其钻探过程中土壤和地下水中明显观察到NAPLs存在时，可将点位所在区域判定为污染源区；",
    )
    add_paragraph(
        doc,
        "④以上条件均不满足，应开展监测，至少3次监测观察到NAPLs存在时，可将点位所在区域判定为污染源区。",
    )
    add_paragraph(doc, "（2）污染羽判定原则")
    add_paragraph(
        doc,
        "①其他辅助条件满足三个及以上，可直接依据微扰动调查成果，将点位所在区域判定为污染羽；",
    )
    add_paragraph(
        doc,
        "②其他辅助条件满足两个及以下，须结合采样检测进行验证，当土壤或地下水中有污染物检出时，可将点位所在区域判定为污染羽。",
    )

    all_image_paths = [
        "./ref/KMEANS-氡气.png",
        "./ref/KMEANS-VOCs.png",
        "./ref/KMEANS-CH4.png",
        "./ref/KMEANS-CO2.png",
        "./ref/KMEANS-O2.png",
        "./ref/KMEANS-功能基因.png",
    ]
    image_paths = []
    for path in all_image_paths:
        if Path(path).exists():
            image_paths.append(path)
    # 计算表格的行数
    print(image_paths)
    rows = (len(image_paths) + 2 - 1) // 2

    # 创建表格
    table = doc.add_table(rows=rows, cols=2)

    # 填充表格单元格
    for i, image_path in enumerate(image_paths):
        row = i // 2
        col = i % 2
        cell = table.cell(row, col)
        if Path(image_path).exists():
            cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(2))
        else:
            cell.add_paragraph()
    add_pic_header(doc, "图5-1 K-MEANS法辅助确定环境背景值")
    add_paragraph(
        doc,
        "按照上述确定的环境背景和明显异常的阈值，将VOCs、CO2、CH4大于阈值的确定为异常点位，小于阈值的确定为无异常点位；O2小于阈值的确定为异常点位，大于阈值的确定为无异常点位，点位分布如图所示。",
    )
    columns = [
        "氡气异常低",
        "VOCs异常高",
        "CO2异常高",
        "O2异常低",
        "CH4异常高",
        "功能基因异常高",
    ]
    labels = ["氡气", "VOCs", "CO2", "O2", "CH4", "功能基因"]
    all_image_paths = [
        "./pic/氡气.png",
        "./pic/VOCs.png",
        "./pic/CH4.png",
        "./pic/CO2.png",
        "./pic/O2.png",
        "./pic/功能基因.png",
    ]
    image_paths = []
    for path in all_image_paths:
        if Path(path).exists():
            image_paths.append(path)
    # 计算表格的行数
    print(image_paths)
    rows = (len(image_paths) + 2 - 1) // 2

    # 创建表格
    table = doc.add_table(rows=rows, cols=2)

    # 填充表格单元格
    for i, image_path in enumerate(image_paths):
        row = i // 2
        col = i % 2
        cell = table.cell(row, col)
        if Path(image_path).exists():
            cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(2))
        else:
            cell.add_paragraph()
    add_pic_header(doc, "Fig5-2 微扰动调查异常点位分布图")
    doc.add_paragraph("")
    return doc


if __name__ == "__main__":
    file = auto_report_EN()
    file.save("C:\\Users\\apple\\Desktop\\auto_report.docx")
