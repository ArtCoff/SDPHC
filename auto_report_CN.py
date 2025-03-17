# from Method_Functions import 绘制点位分布, 绘制污染点位分布V2, 绘制保存异常点位
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from datetime import datetime
from docx.oxml.ns import qn  # 用于设置中文字体
from pathlib import Path
import geopandas as gpd
from docx.shared import RGBColor


def set_heading_style(heading, level=1):
    """
    设置一级标题样式：
    - 字体为黑体
    - 字号为三号
    - 加粗
    - 顶格（左对齐）

    :heading: 段落对象
    :text: 标题文本
    """
    # 设置标题的字体和字号
    run = heading.runs[0]  # 获取标题的第一个 Run 对象
    if level == 1:
        run.font.size = Pt(16)  # 设置字号为三号字（16磅）
        run.bold = True  # 设置加粗
        run.font.name = "SimHei"  # 设置字体为黑体
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "黑体")  # 设置中文为黑体
        run.font.color.rgb = RGBColor(0, 0, 0)

        # 设置段落格式（顶格、左对齐）
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
    elif level == 2:
        run.font.size = Pt(12)  #
        run.bold = True  # 设置加粗
        run.font.name = "SimSong"  # 设置字体为黑体
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 设置中文为黑体
        run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置段落格式（顶格、左对齐）
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐


def set_paragraph_style(
    paragraph,
    is_chinese=True,
    font_size=12,
    line_spacing=1.5,
    space_before=6,
    space_after=6,
):
    """
    设置段落的字体、字号、行距和段落间距
    :param paragraph: 段落对象
    :param is_chinese: 是否是中文，True 表示中文，False 表示英文
    :param font_size: 字号（磅）
    :param line_spacing: 行距
    :param space_before: 段前间距
    :param space_after: 段后间距
    """
    # 设置首行缩进 2 字符

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Pt(24)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 设置两端对齐
    paragraph_format.line_spacing = line_spacing  # 设置行距
    paragraph_format.space_before = Pt(space_before)  # 段前间距
    paragraph_format.space_after = Pt(space_after)  # 段后间距

    # 设置 Run 的字体属性
    for run in paragraph.runs:
        if is_chinese:
            run.font.name = "Times New Roman"  # 中文设置
            run.font.color.rgb = RGBColor(0, 0, 0)
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 设置黑体
        else:
            run.font.name = "Times New Roman"  # 英文设置
            run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(font_size)  # 设置字号


def add_paragraph(doc, text):
    """
    插入自定义格式的段落：小四字号、宋体、1.5倍行距、两端对齐、首行缩进2字符。

    :param doc: Document 对象
    :param text: 要插入的文本
    """
    # 插入段落
    paragraph = doc.add_paragraph(text)

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.first_line_indent = Pt(
        24
    )  # 首行缩进2字符（每个字符12磅，2字符即24磅）

    # 设置字体样式
    for run in paragraph.runs:
        run.font.name = "Times New Roman"  # 设置字体为宋体
        run.font.size = Pt(12)  # 小四字号（12磅）
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 确保设置为宋体

    return paragraph


def add_table_header(doc, text):
    # 插入段落
    paragraph = doc.add_paragraph(text)

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.space_before = Pt(6)  # 段前间距
    paragraph_format.space_after = Pt(0)  # 段后间距
    # 设置字体样式
    for run in paragraph.runs:
        run.font.name = "Times New Roman"  # 设置字体为宋体
        run.font.size = Pt(12)  # 小四字号（12磅）
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 确保设置为宋体

    return paragraph


def add_pic_header(doc, text):
    # 插入段落
    paragraph = doc.add_paragraph(text)

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.space_before = Pt(0)  # 段前间距
    paragraph_format.space_after = Pt(6)  # 段后间距
    # 设置字体样式
    for run in paragraph.runs:
        run.font.name = "Times New Roman"  # 设置字体为宋体
        run.font.size = Pt(12)  # 小四字号（12磅）
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 确保设置为宋体

    return paragraph


from PIL import Image
from docx import Document
from docx.shared import Inches


def insert_image(doc, image_path, width=None, height=None):
    """
    插入图片，并支持设置宽度和高度，图片默认居中显示。

    :param doc: Document 对象
    :param image_path: 图片路径
    :param width: 图片宽度（可选，单位为英寸，默认不设置）
    :param height: 图片高度（可选，单位为英寸，默认不设置）
    """
    if not Path(image_path).exists():
        print(f"图片路径无效：{image_path}")
        return
    # 创建段落
    paragraph = doc.add_paragraph()

    # 在段落中插入图片
    run = paragraph.add_run()

    # 设置图片宽度和高度
    if width and height:
        run.add_picture(
            image_path, width=Inches(width), height=Inches(height)
        )  # 设置宽高
    elif width:
        run.add_picture(image_path, width=Inches(width))  # 只设置宽度
    elif height:
        run.add_picture(image_path, height=Inches(height))  # 只设置高度
    else:
        run.add_picture(image_path)  # 默认不设置，按原图大小插入

    # 设置图片所在段落居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中显示图片


def apply_heading_style_to_all_titles(doc):
    """
    遍历文档中的所有一级标题，应用标题样式。

    :param doc: Document 对象
    """
    for paragraph in doc.paragraphs:
        # 判断该段落是否为一级标题
        if paragraph.style.name == "Heading 1":  # 默认为一级标题的样式名为 "Heading 1"
            set_heading_style(paragraph, level=1)  # 设置标题样式
        elif paragraph.style.name == "Heading 2":
            set_heading_style(paragraph, level=2)
        # elif paragraph.style.name == "Normal":
        #     set_paragraph_style(paragraph)


def auto_report(gdf, background_file=None, final_boundarys=None):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION
    from datetime import datetime
    from docx.oxml.ns import qn  # 用于设置中文字体
    from pathlib import Path
    import geopandas as gpd

    return
    # 创建文档
    doc = Document()

    # 封面标题

    # 添加空段落，使标题位于页面中部
    for _ in range(8):  # 根据页面长度适当调整数量
        i = doc.add_paragraph()
        i.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = doc.add_paragraph()
    title_run = title.add_run("微扰动污染调查报告")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    title_run.font.name = "SimHei"  # 设置中文字体
    title_run.font.size = Pt(40)  # 设置字体大小为三号
    title_run.bold = True  # 加粗
    r = title_run._element  # 处理中文字体
    r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 添加单位名称
    # 添加空段落，使标题位于页面中部
    for _ in range(12):  # 根据页面长度适当调整数量
        i = doc.add_paragraph()
        i.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    organization_paragraph = doc.add_paragraph()
    organization_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    organization_run = organization_paragraph.add_run(
        "中国地质科学院水文地质环境地质研究所"
    )
    organization_run.font.name = "SimHei"
    organization_run.font.size = Pt(16)
    r = organization_run._element  # 处理中文字体
    r.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # 添加时间
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    date_run = date_paragraph.add_run(datetime.now().strftime("%Y年 %m月"))
    date_run.font.name = "SimHei"
    date_run.font.size = Pt(16)
    r = date_run._element  # 处理中文字体
    r.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # 添加分节符
    doc.add_section(WD_SECTION.NEW_PAGE)

    # 添加正文示例
    doc.add_heading("1 调查目的", level=1)
    add_paragraph(
        doc,
        "基于XX石化自行监测结果，针对在产企业作业空间受限、地下管线密布、对防火、防爆等安全要求极高的特点，微扰动调查技术可减少动土作业的安全隐患，为在产企业提供“边生产，边防控”双重责任履行的方式和手段，具有安全、高效、便捷、快速、多维度、高精度、低成本等优点，本项目中采用微扰动技术识别场地污染源区，确定厂区污染范围，为污染隐患消除及后续风险管控方向提供技术支撑。",
    )

    add_paragraph(
        doc,
        "微扰动调查分为两个阶段，第一阶段为已有信息和调查方式确认，第二阶段为污染源区及污染范围识别。",
    )

    doc.add_heading("2 已有信息和调查方式确认", level=1)
    doc.add_heading("2.1 工作内容", level=2)
    add_paragraph(
        doc,
        "包括现场踏勘、人员访谈、微扰动调查、数据评估与结果分析等工作，确认场地调查方式。",
    )
    doc.add_heading("2.2 现场踏勘和人员访谈", level=2)
    add_paragraph(
        doc,
        "调查人员对企业自行监测等资料进行分析，包括对监测资料的可靠性和合理性的分析评估。对企业进行了现场踏勘和人员访谈，在企业安全环保负责人的引导下，针对土壤污染隐患排查识别的重点场所和重点设施设备及其周边区域，识别了易燃、易爆、易中毒、易高温、易发生机械伤害等高危险区域及受限空间，观察和记录了地下管线、电缆、隐蔽设施等构筑物的分布情况，并梳理开展了调查所需条件清单（示例表2-1），综合成本、时效需求，与厂方确认调查方式。",
    )

    doc.add_heading("3 微扰动调查", level=1)
    doc.add_heading("3.1 微扰动调查原理", level=2)
    add_paragraph(doc, "（1）表层土壤气体调查技术原理")
    add_paragraph(
        doc,
        "挥发性有机物（VOCs）在土壤和地下水迁移过程中，伴随着向地表的蒸汽挥发和生物降解两过程（图3-1）。以研究以较为成熟的石油烃类为例，烃类的蒸汽挥发和生物降解过程交互，呈现矿化、甲烷化、甲烷氧化等垂向氧化还原分带现象。此过程中产生的CO2、CH4、H2S、N2等气体会向地表传输，同时由空气向地下输入的O2被消耗，最终导致表层土壤中气体组成异常。因此，通过微扰动调查技术检测表层土壤气体组成，可指示土壤和地下水VOCs污染情况。",
    )
    insert_image(doc, "./ref/微扰动原理.png", width=4)
    add_pic_header(doc, "图3-1 微扰动调查技术原理图")
    doc.add_paragraph()
    add_paragraph(doc, "（2）氡评估NAPL技术原理")
    add_paragraph(
        doc,
        "氡技术评估NAPL污染是基于氡易溶于有机溶剂而导致土壤孔隙中氡浓度的变化来对NAPL污染进行定位及量调查。土壤中不存在NAPL 时，由镭衰变产生的自由氡通过扩散与运移向土壤表面迁移而进入到大气中，当土壤中存在 NAPL 时，由于氡优先溶解于有机溶剂中，会导致土壤孔隙氡浓度大幅下降。应用氡示踪剂技术在评估土壤及地下水 NAPL 污染时，通过测量土壤受 NAPL 污染前后孔隙氡浓度的空间分布，可以迅速而准确的对 NAPL 污染源进行定位及定量分析土壤地下水中NAPL污染饱和度。",
    )
    doc.add_heading("3.2 调查方法", level=2)
    add_paragraph(
        doc,
        "选用土壤深度0.7m以浅动土作业的氡气和VOCs、CO2、O2、CH4等微扰动调查技术来确定场地调查方式。",
    )
    add_paragraph(
        doc, "调查技术包含前期准备、钻孔、连接仪器装置、取气检测等4个主要步骤，具体为："
    )
    add_paragraph(doc, "（1）前期准备")
    add_paragraph(
        doc,
        "调查前准备相关仪器设备，包括：钻入设备（铜锤、铜钎）、取气装置及聚四氟乙烯连接管、泵吸式多参数气体检测仪（包含VOCs、CO2、O2、CH4、NH4及HCHO气体）、泵吸式测氡仪。",
    )
    add_paragraph(doc, "（2）钻孔")
    add_paragraph(
        doc,
        "利用铜锤、铜钎防爆钻入设备，在调查点位，开凿深约0.7m、直径约2cm的钻孔。每个孔位钻入前征询企业相关人员的意见。",
    )
    add_paragraph(doc, "（3）连接仪器装置")
    add_paragraph(
        doc,
        "将带有内径约2mm的聚四氟乙烯管的取气装置（或塞子）密封钻孔孔口，防止取气过程中空气的进入；聚四氟乙烯导气管连接取气装置的出气口和气体检测仪的进气口。连接时检查各进出气口及管路，以管路防堵塞或漏气。",
    )
    add_paragraph(doc, "（4）取气检测")
    add_paragraph(
        doc,
        "打开泵吸式检测仪，取气与检测同步进行。检测过程中，实时记录仪器读数，待仪器各指标读数均稳定或出现拐点后，取气检测工作结束，选取各指标的稳定值或拐点值作为该指标的最终测试结果。",
    )
    add_paragraph(
        doc,
        "氡气测定时在其他气体检测点位旁边进行单独钻孔进行测试。氡浓度使用RAD7测氡仪（图3-2）进行测量，利用α能谱测定法的原理，来对空气中的氡含量进行测定，取气装置示意图如图3-3所示。",
    )
    insert_image(doc, "./ref/RAD-7测氡仪.png", width=4)
    add_pic_header(doc, "图3-2 RAD-7测氡仪")
    insert_image(doc, "./ref/取气装置.png", width=2.6)
    add_pic_header(doc, "图3-3 取气装置")
    doc.add_paragraph("")
    doc.add_heading("4 调查点位", level=1)
    add_paragraph(
        doc,
        f"调查点位设置在厂区内，包括厂区内部及周边区域，共计{gdf.shape[0]}个点位。调查点位设置如下图所示：",
    )
    绘制点位分布(gdf, boundary_polygon_file=background_file)
    绘制污染点位分布V2(gdf, boundary_polygon_file=background_file, save=True)

    insert_image(doc, "./ref/监测点位分布图.png", width=4)
    add_pic_header(doc, "图4-1 调查点位分布图")

    # 填充表格
    header = ["点位", "氡气", "VOCs", "CO2", "O2", "CH4", "功能基因"]
    table_header = [
        "点位",
        "氡气(Bq/m³)",
        "VOCs(ppb)",
        "CO2(ppm)",
        "O2(%)",
        "CH4(%)",
        "功能基因(copies/g)",
    ]
    selected_data = gdf[header]
    rows, cols = selected_data.shape
    print(rows, cols)
    add_table_header(doc, "表4-1 微扰动调查点位数据")
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = "Table Grid"
    for i, col_name in enumerate(table_header):
        table.cell(0, i).text = col_name
    for row_idx, row_data in selected_data.iterrows():
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx + 1, col_idx).text = str(value)
    doc.add_heading("5 数据评估与结果分析", level=1)
    doc.add_heading("5.1 判定原则", level=2)
    add_paragraph(
        doc,
        "调查人员须具备一定专业知识和工作经验，推荐综合所有指标，以K-MEANS聚类法、参考值法等，识别环境背景区域，确定微扰动调查指标的环境环境背景值。与环境环境背景值比较，其中氡气明显异常低，是微扰动调查判定污染源区的必要条件。以如下其他四个辅助条件圈定污染羽，并可辅以判定污染源区：①VOCs明显异常高；②CH4明显异常高；③CO2明显异常高、O2明显异常低；④功能基因明显异常高。",
    )
    add_paragraph(doc, "（1）污染源区判定原则")
    add_paragraph(
        doc,
        "①必要条件满足，且其他四个辅助条件均满足，可直接依据微扰动调查成果，将点位所在区域判定为污染源区；",
    )
    add_paragraph(
        doc,
        "②必要条件满足，且其他辅助条件满足三个，可结合采样检测进行验证，当表层土壤和地下水中污染物浓度含量均超过相关标准限值要求时，将点位所在区域判定为污染源区；",
    )
    add_paragraph(
        doc,
        "③必要条件满足，且其他辅助条件满足两个及以下，须结合采样检测进行验证，尤其钻探过程中土壤和地下水中明显观察到NAPLs存在时，可将点位所在区域判定为污染源区；",
    )
    add_paragraph(
        doc,
        "④以上条件均不满足，应开展监测，至少3次监测观察到NAPLs存在时，可将点位所在区域判定为污染源区。",
    )
    add_paragraph(doc, "（2）污染羽判定原则")
    add_paragraph(
        doc,
        "①其他辅助条件满足三个及以上，可直接依据微扰动调查成果，将点位所在区域判定为污染羽；",
    )
    add_paragraph(
        doc,
        "②其他辅助条件满足两个及以下，须结合采样检测进行验证，当土壤或地下水中有污染物检出时，可将点位所在区域判定为污染羽。",
    )

    doc.add_heading("5.2 污染点位判定", level=2)
    add_paragraph(doc, "(1) 异常点位识别")
    add_paragraph(
        doc,
        f"在厂区共布设了微扰动调查点位{rows}个，综合考虑氡气、VOCs、CO2、O2、CH4、流场、污染物及生产工艺等，采用K-MEANS聚类法和参照值法，确定了①VOCs的环境背景和明显异常的阈值为{final_boundarys[1]}ppb，②CH4的环境背景和明显异常的阈值为{final_boundarys[4]}%，③CO2的环境背景和明显异常的阈值为{final_boundarys[2]}ppm，O2的环境背景和明显异常的阈值为{final_boundarys[3]}%。",
    )
    all_image_paths = [
        "./ref/KMEANS-氡气.png",
        "./ref/KMEANS-VOCs.png",
        "./ref/KMEANS-CH4.png",
        "./ref/KMEANS-CO2.png",
        "./ref/KMEANS-O2.png",
        "./ref/KMEANS-功能基因.png",
    ]
    image_paths = []
    for path in all_image_paths:
        if Path(path).exists():
            image_paths.append(path)
    # 计算表格的行数
    print(image_paths)
    rows = (len(image_paths) + 2 - 1) // 2

    # 创建表格
    table = doc.add_table(rows=rows, cols=2)

    # 填充表格单元格
    for i, image_path in enumerate(image_paths):
        row = i // 2
        col = i % 2
        cell = table.cell(row, col)
        if Path(image_path).exists():
            cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(2))
        else:
            cell.add_paragraph()
    add_pic_header(doc, "图5-1 K-MEANS法辅助确定环境背景值")
    add_paragraph(
        doc,
        "按照上述确定的环境背景和明显异常的阈值，将VOCs、CO2、CH4大于阈值的确定为异常点位，小于阈值的确定为无异常点位；O2小于阈值的确定为异常点位，大于阈值的确定为无异常点位，点位分布如图所示。",
    )
    columns = [
        "氡气异常低",
        "VOCs异常高",
        "CO2异常高",
        "O2异常低",
        "CH4异常高",
        "功能基因异常高",
    ]
    labels = ["氡气", "VOCs", "CO2", "O2", "CH4", "功能基因"]
    for index, value in enumerate(columns):
        绘制保存异常点位(gdf=gdf, column=value, label=labels[index])
    all_image_paths = [
        "./pic/氡气.png",
        "./pic/VOCs.png",
        "./pic/CH4.png",
        "./pic/CO2.png",
        "./pic/O2.png",
        "./pic/功能基因.png",
    ]
    image_paths = []
    for path in all_image_paths:
        if Path(path).exists():
            image_paths.append(path)
    # 计算表格的行数
    print(image_paths)
    rows = (len(image_paths) + 2 - 1) // 2

    # 创建表格
    table = doc.add_table(rows=rows, cols=2)

    # 填充表格单元格
    for i, image_path in enumerate(image_paths):
        row = i // 2
        col = i % 2
        cell = table.cell(row, col)
        if Path(image_path).exists():
            cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(2))
        else:
            cell.add_paragraph()
    add_pic_header(doc, "图5-2 微扰动调查异常点位分布图")
    add_paragraph(doc, "(2) 污染点位判定")
    doc.add_paragraph("")

    doc.add_paragraph("")
    add_table_header(doc, "表5-1 污染点位判定矩阵")
    abnormal_header = [
        "点位",
        "氡气异常低",
        "VOCs异常高",
        "CO2异常高",
        "O2异常低",
        "CH4异常高",
        "功能基因异常高",
    ]
    selected_data = gdf[abnormal_header]
    rows, cols = selected_data.shape
    abnormal_header_table = doc.add_table(rows=rows + 1, cols=cols)
    abnormal_header_table.style = "Table Grid"
    for i, col_name in enumerate(abnormal_header):
        abnormal_header_table.cell(0, i).text = col_name
    for row_idx, row_data in selected_data.iterrows():
        for col_idx, value in enumerate(row_data):
            abnormal_header_table.cell(row_idx + 1, col_idx).text = str(value)
    insert_image(doc, "./ref/污染点位分布图.png", width=4)
    add_table_header(doc, "图5-2 场区污染点位分布图")

    apply_heading_style_to_all_titles(doc)
    return doc


# if __name__ == "__main__":
#     import geopandas as gpd

#     gdf = gpd.read_file("./Result.gpkg")
#     file = auto_report(gdf, None)
#     file.save("C:\\Users\\g2382\\Desktop\\auto_report.docx")
