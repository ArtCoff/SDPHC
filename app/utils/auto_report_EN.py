import os
import psutil
from datetime import datetime
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn


def is_file_locked(file_path):
    """检测文件是否被其他进程占用"""
    try:
        # 尝试以只读模式打开文件
        with open(file_path, "r"):
            pass
        return False
    except IOError:
        return True


def find_and_kill_file_handles(file_path):
    """查找并终止占用文件的进程"""
    for proc in psutil.process_iter(["pid", "name", "open_files"]):
        try:
            if proc.info["open_files"]:
                for file in proc.info["open_files"]:
                    if file.path == os.path.abspath(file_path):
                        print(f"发现占用进程: {proc.info['name']} (PID: {proc.pid})")
                        try:
                            p = psutil.Process(proc.pid)
                            p.terminate()
                            p.wait(timeout=3)
                            print(f"已终止进程 PID {proc.pid}")
                        except psutil.NoSuchProcess:
                            print(f"进程 {proc.pid} 已不存在")
                        except Exception as e:
                            print(f"终止进程失败: {e}")
        except psutil.NoSuchProcess:
            continue


def save_docx_safely(doc, file_path):
    """安全保存 docx 文件，自动处理文件占用问题"""
    if os.path.exists(file_path):
        if is_file_locked(file_path):
            print(f"文件 {file_path} 正在被占用，尝试关闭占用进程...")
            find_and_kill_file_handles(file_path)

    # 保存文件
    try:
        doc.save(file_path)
        print(f"文件已成功保存至: {file_path}")
    except Exception as e:
        print(f"保存文件失败: {e}")


def set_heading_style(heading, level=1):
    # 设置标题的字体和字号
    run = heading.runs[0]
    if level == 1:
        run.font.size = Pt(16)
        run.bold = True  # 设置加粗
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置段落格式（顶格、左对齐）
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif level == 2:
        run.font.size = Pt(12)  #
        run.bold = True  # 设置加粗
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置段落格式（顶格、左对齐）
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif level == 3:
        run.font.size = Pt(12)  #
        run.bold = False  # 设置加粗
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置段落格式（顶格、左对齐）
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def set_paragraph_style(
    paragraph,
    is_chinese=False,
    font_size=12,
    line_spacing=1.5,
    space_before=6,
    space_after=6,
):
    """
    设置段落的字体、字号、行距和段落间距
    :param paragraph: 段落对象
    :param is_chinese: 是否是中文，True 表示中文，False 表示英文
    :param font_size: 字号（磅）
    :param line_spacing: 行距
    :param space_before: 段前间距
    :param space_after: 段后间距
    """
    # 设置首行缩进 2 字符

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Pt(24)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 设置两端对齐
    paragraph_format.line_spacing = line_spacing  # 设置行距
    paragraph_format.space_before = Pt(space_before)  # 段前间距
    paragraph_format.space_after = Pt(space_after)  # 段后间距

    # 设置 Run 的字体属性
    for run in paragraph.runs:
        if is_chinese:
            run.font.name = "simsun"  # 中文设置
            run.font.color.rgb = RGBColor(0, 0, 0)
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 设置黑体
        else:
            run.font.name = "Times New Roman"  # 英文设置
            run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(font_size)  # 设置字号


def add_table_header(doc, text):
    # 插入段落
    paragraph = doc.add_paragraph(text)

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.space_before = Pt(6)  # 段前间距
    paragraph_format.space_after = Pt(0)  # 段后间距
    # 设置字体样式
    for run in paragraph.runs:
        run.font.name = "Times New Roman"  # 设置字体为宋体
        run.font.size = Pt(12)  # 小四字号（12磅）
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 确保设置为宋体

    return paragraph


def add_pic_header(doc, text):
    # 插入段落
    paragraph = doc.add_paragraph(text)

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.space_before = Pt(0)  # 段前间距
    paragraph_format.space_after = Pt(6)  # 段后间距
    # 设置字体样式
    for run in paragraph.runs:
        run.font.name = "Times New Roman"  # 设置字体为宋体
        run.font.size = Pt(12)  # 小四字号（12磅）
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 确保设置为宋体

    return paragraph


def insert_image(doc, image_path, width=None, height=None):
    """
    插入图片，并支持设置宽度和高度，图片默认居中显示。

    :param doc: Document 对象
    :param image_path: 图片路径
    :param width: 图片宽度（可选，单位为英寸，默认不设置）
    :param height: 图片高度（可选，单位为英寸，默认不设置）
    """
    if not Path(image_path).exists():
        print(f"图片路径无效：{image_path}")
        return
    # 创建段落
    paragraph = doc.add_paragraph()

    # 在段落中插入图片
    run = paragraph.add_run()

    # 设置图片宽度和高度
    if width and height:
        run.add_picture(
            image_path, width=Inches(width), height=Inches(height)
        )  # 设置宽高
    elif width:
        run.add_picture(image_path, width=Inches(width))  # 只设置宽度
    elif height:
        run.add_picture(image_path, height=Inches(height))  # 只设置高度
    else:
        run.add_picture(image_path)  # 默认不设置，按原图大小插入

    # 设置图片所在段落居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中显示图片


def setup_styles(doc):
    # =====================================
    # 正文样式 (Normal)
    # =====================================
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"  # 英文字体
    normal_style.font.size = Pt(12)  # 字号
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    # 段落格式
    normal_style.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进
    normal_style.paragraph_format.line_spacing = 1.5  # 行距
    normal_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐

    # =====================================
    # List样式 (List Paragraph)
    list_style = doc.styles["List"]
    list_style.font.name = "Times New Roman"  # 英文字体
    list_style.font.size = Pt(11)  # 字号
    list_style.font.color.rgb = RGBColor(0, 0, 0)
    # 段落格式
    list_style.paragraph_format.first_line_indent = Cm(0)  # 首行缩进
    list_style.paragraph_format.line_spacing = 1.15
    list_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # =====================================
    # 标题样式 (Heading 1-3)
    # =====================================
    heading_configs = {
        "Heading 1": {
            "font": "Arial Black",
            "size": Pt(16),
            "bold": True,
            "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "space_before": Pt(24),
            "space_after": Pt(12),
        },
        "Heading 2": {
            "font": "Arial Black",
            "size": Pt(14),
            "bold": True,
            "italic": True,
            "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "space_before": Pt(18),
            "space_after": Pt(6),
        },
        "Heading 3": {
            "font": "Times New Roman",
            "size": Pt(12),
            "bold": False,
            "italic": True,
            "underline": False,
            "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "space_before": Pt(12),
            "space_after": Pt(3),
        },
    }

    for style_name, config in heading_configs.items():
        style = doc.styles[style_name]
        # 字体设置
        style.font.name = config["font"]
        style.font.size = config["size"]
        style.font.bold = config.get("bold", False)
        style.font.italic = config.get("italic", False)
        style.font.underline = config.get("underline", False)
        style.font.color.rgb = RGBColor(0, 0, 0)  # 统一标题颜色
        # 段落格式
        style.paragraph_format.alignment = config["alignment"]
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = config["space_before"]
        style.paragraph_format.space_after = config["space_after"]
        style.paragraph_format.first_line_indent = Pt(0)  # 首行缩进 0 字符


def add_cover(doc):
    for _ in range(6):  # 根据页面长度适当调整数量
        i = doc.add_paragraph()
        i.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = doc.add_paragraph()
    title_run = title.add_run(
        "XXX Petroleum Hydrocarbon Contaminated Site Investigation Report"
    )
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    title_run.font.name = "Times New Roman"  # 设置中文字
    title_run.font.size = Pt(20)  # 设置字体大小为三号
    title_run.bold = True  # 加粗

    # 添加空段落，使标题位于页面中部
    for _ in range(10):  # 根据页面长度适当调整数量
        i = doc.add_paragraph()
        i.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    organization_paragraph = doc.add_paragraph()
    organization_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    organization_run = organization_paragraph.add_run(
        "XX environmental technology company"
    )
    organization_run.font.name = "Times New Roman"
    organization_run.font.size = Pt(12)

    # 添加时间
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    date_run = date_paragraph.add_run(datetime.now().strftime("%Y-%m"))
    date_run.font.name = "Times New Roman"
    date_run.font.size = Pt(16)
    # 添加分节符
    doc.add_section(WD_SECTION.NEW_PAGE)
    return doc


def auto_report_EN():
    # 创建文档
    doc = Document()
    # 设置默认样式
    setup_styles(doc)
    doc = add_cover(doc)  # 添加封面

    doc.add_heading("1 Guidelines for automated report generation", level=1)
    doc.add_heading("1.1 Code writing", level=3)
    doc.add_paragraph(
        "The code is written in Python, and the report generation is based on the python-docx library. The code is modularized to facilitate maintenance and reuse. The code of Auto_report is divided into three parts: data processing, report generation, and report export.",
    )
    doc.add_paragraph(
        "The data processing part is responsible for reading and processing the data, and generating the required figures.The auto_report_cache folder exists in the installation directory and is used to temporarily save finished drawings or processed data.",
    )
    doc.add_paragraph(
        "The report generation part is responsible for generating the report template and inserting the processed data and images.The code for this section is saved as a separate python file that can be called from the software GUI page and new templates are constantly being written for specific needs.",
    )
    doc.add_paragraph(
        "The report export part is responsible for exporting the generated report to a Word document.Usually, the code for the report export section is fixed and only requires a call to the System Resource Manager in the software.",
    )
    doc.add_heading("1.2 Basic format control", level=2)
    doc.add_paragraph(
        "The report is generated in Word format and the basic format is controlled by the python-docx library.For example, the body of the font Times New Roman, font size of 12 pounds, color black, line spacing is 1.5 times, the first line indented 2 characters, alignment is aligned at both ends. The way to control by code is:"
    )
    doc.add_paragraph(
        "For example, the body of the font Times New Roman, font size of 12 pounds, color black, line spacing is 1.5 times, the first line indented 2 characters, alignment is aligned at both ends. The way to control by code is:"
    )

    code = """normal_style = doc.styles["Normal"]\nnormal_style.font.name = "Times New Roman"\nnormal_style.font.size = Pt(12)\nnormal_style.font.color.rgb = RGBColor(0, 0, 0)\nnormal_style.paragraph_format.first_line_indent = Cm(0.74)\nnormal_style.paragraph_format.line_spacing = 1.5\nnormal_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Courier New"  # 设置字体为等宽字体
    run.font.size = Pt(10)  # 设置字号为10磅
    run.font.color.rgb = RGBColor(0, 0, 0)  #
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 段落级设置
    paragraph.paragraph_format.first_line_indent = Pt(0)  # 段落级设置
    paragraph.paragraph_format.line_spacing = 1  # 段落级设置
    # 填充表格
    doc.add_heading("1.3 Table and figure generation", level=2)
    doc.add_paragraph(
        "The table and figure generation is based on the add_table and add_picture functions, which provide a simple and easy-to-use interface for generating tables and figures.",
    )
    doc.add_paragraph(
        "However, the python-docx library has some limitations on the insertion of tables and images: it only provides the most direct way of insertion, and personalization requires the writing of more complex code that needs to manipulate the underlying XML.",
    )
    doc.add_paragraph(
        "The following is an example of inserting a customized table (the data used are randomly generated and have no practical significance):"
    )

    import random

    table_header = [
        "ID",
        "Rn(Bq/m³)",
        "VOCs(ppb)",
        "CO2(ppm)",
        "O2(%)",
        "CH4(%)",
        "FG(copies/g)",
    ]
    add_table_header(doc, "Table1-1 Non-invasive survey site data")
    table = doc.add_table(rows=6, cols=len(table_header))
    table.style = "Table Grid"
    for i, col_name in enumerate(table_header):
        table.cell(0, i).text = col_name
        for j in range(1, 6):
            table.cell(j, i).text = str(random.randint(10, 100))

    # set picture
    doc.add_paragraph()
    doc.add_paragraph("The following is an example of inserting a group of pictures:")

    all_image_paths = [
        "./auto_report_cache/VOCs.png",
        "./auto_report_cache/CH4.png",
        "./auto_report_cache/CO2.png",
        "./auto_report_cache/O2.png",
        "./auto_report_cache/FG.png",
    ]
    image_paths = []
    for path in all_image_paths:
        if Path(path).exists():
            image_paths.append(path)
    # 计算表格的行数
    print(image_paths)
    rows = (len(image_paths) + 2 - 1) // 2

    # 创建表格
    table = doc.add_table(rows=rows, cols=2)

    # 填充表格单元格
    for i, image_path in enumerate(image_paths):
        row = i // 2
        col = i % 2
        cell = table.cell(row, col)
        if Path(image_path).exists():
            cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(2))
        else:
            cell.add_paragraph()
    add_pic_header(doc, "Figure 1-1 Example of inserting a group of pictures")
    # 应用样式
    return doc


if __name__ == "__main__":
    file = auto_report_EN()
    file.save("auto_report.docx")
