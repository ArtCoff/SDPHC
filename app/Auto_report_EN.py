from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn  # 用于设置中文字体

# 其他库引用
from datetime import datetime
from pathlib import Path
import geopandas as gpd
from PIL import Image

# 内部函数引用


# def setup_default_styles(doc):
#     # ================================
#     # 设置默认正文样式（Normal）
#     # ================================
#     normal_style = doc.styles["Normal"]
#     # 字体设置
#     normal_style.font.name = "Times New Roman"  # 西文字体
#     normal_style.font.size = Pt(12)  # 小四号字
#     normal_style.font.color.rgb = RGBColor(0, 0, 0)
#     # 中文字体（需操作底层 XML）
#     rpr = normal_style.element.get_or_add_rPr()
#     rpr.get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

#     # 段落格式设置
#     paragraph_format = normal_style.paragraph_format
#     paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进 2字符
#     paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE  # 行距类型
#     paragraph_format.line_spacing = 1.5  # 1.5倍行距
#     paragraph_format.space_before = Pt(0)  # 段前间距
#     paragraph_format.space_after = Pt(0)  # 段后间距
#     paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐

#     # ================================
#     # 设置标题样式（支持 3 级标题）
#     # ================================
#     heading_configs = {
#         "Heading 1": {
#             "size": Pt(16),  # 三号字
#             "bold": True,
#             "font_en": "Times New Roman",
#             "font_cn": "黑体",
#             "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,  # 居中对齐
#             "line_spacing": 1.5,
#             "space_before": Pt(24),
#             "space_after": Pt(12),
#         },
#         "Heading 2": {
#             "size": Pt(15),  # 小三号
#             "bold": True,
#             "font_en": "Times New Roman",
#             "font_cn": "宋体",
#             "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,  # 左对齐
#             "line_spacing": 1.0,
#             "space_before": Pt(18),
#             "space_after": Pt(6),
#         },
#         "Heading 3": {
#             "size": Pt(14),  # 四号字
#             "bold": True,
#             "font_en": "Times New Roman",
#             "font_cn": "宋体",
#             "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,  # 右对齐
#             "line_spacing": WD_LINE_SPACING.EXACTLY,  # 固定值
#             "line_spacing_pt": Pt(20),  # 固定 20磅
#             "space_before": Pt(12),
#             "space_after": Pt(3),
#         },
#     }

#     for style_name, config in heading_configs.items():
#         heading_style = doc.styles[style_name]
#         # 字体设置
#         heading_style.font.size = config["size"]
#         heading_style.font.bold = config["bold"]
#         heading_style.font.name = config["font_en"]
#         # 中文字体
#         rpr = heading_style.element.get_or_add_rPr()
#         rpr.get_or_add_rFonts().set(qn("w:eastAsia"), config["font_cn"])
#         # 段落格式
#         pf = heading_style.paragraph_format
#         pf.alignment = config["alignment"]
#         pf.line_spacing_rule = config.get("line_spacing_rule", WD_LINE_SPACING.MULTIPLE)

#         # 支持固定行距（当 line_spacing_pt 存在时）
#         if "line_spacing_pt" in config:
#             pf.line_spacing = config["line_spacing_pt"]
#             pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

#         pf.space_before = config["space_before"]
#         pf.space_after = config["space_after"]
#         # 特殊设置（如大纲级别）
# pf.outline_level = getattr(WD_PARAGRAPH_ALIGNMENT, f"HEADING_{style_name[-1]}")

# ================================
# 新增引用样式（自定义样式）
# ================================
# quote_style = doc.styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)
# quote_style.base_style = doc.styles["Normal"]
# # 字体设置
# quote_style.font.italic = True
# quote_style.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
# # 段落格式
# qpf = quote_style.paragraph_format
# qpf.left_indent = Cm(1.0)  # 左缩进 1厘米
# qpf.right_indent = Cm(1.0)  # 右缩进 1厘米
# qpf.line_spacing = 1.25  # 1.25倍行距
# qpf.shading.background_pattern_color = RGBColor(240, 240, 240)  # 背景色


def set_heading_style(heading, level=1):
    # 设置标题的字体和字号
    run = heading.runs[0]  # 获取标题的第一个 Run 对象
    if level == 1:
        run.font.size = Pt(16)  # 设置字号为三号字（16磅）
        run.bold = True  # 设置加粗
        run.font.name = "Times New Roman"  # 设置字体为黑体
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "黑体")  # 设置中文为黑体
        run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置段落格式（顶格、左对齐）
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
    elif level == 2:
        run.font.size = Pt(12)  #
        run.bold = True  # 设置加粗
        run.font.name = "Times New Roman"  # 设置字体为黑体
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 设置中文为黑体
        run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置段落格式（顶格、左对齐）
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
    elif level == 3:
        run.font.size = Pt(12)  #
        run.bold = False  # 设置加粗
        run.font.name = "Times New Roman"  # 设置字体为黑体
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 设置中文为黑体
        run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置段落格式（顶格、左对齐）
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐


def set_paragraph_style(
    paragraph,
    is_chinese=False,
    font_size=12,
    line_spacing=1.5,
    space_before=6,
    space_after=6,
):
    """
    设置段落的字体、字号、行距和段落间距
    :param paragraph: 段落对象
    :param is_chinese: 是否是中文，True 表示中文，False 表示英文
    :param font_size: 字号（磅）
    :param line_spacing: 行距
    :param space_before: 段前间距
    :param space_after: 段后间距
    """
    # 设置首行缩进 2 字符

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Pt(24)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 设置两端对齐
    paragraph_format.line_spacing = line_spacing  # 设置行距
    paragraph_format.space_before = Pt(space_before)  # 段前间距
    paragraph_format.space_after = Pt(space_after)  # 段后间距

    # 设置 Run 的字体属性
    for run in paragraph.runs:
        if is_chinese:
            run.font.name = "simsun"  # 中文设置
            run.font.color.rgb = RGBColor(0, 0, 0)
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 设置黑体
        else:
            run.font.name = "Times New Roman"  # 英文设置
            run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(font_size)  # 设置字号


def add_table_header(doc, text):
    # 插入段落
    paragraph = doc.add_paragraph(text)

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.space_before = Pt(6)  # 段前间距
    paragraph_format.space_after = Pt(0)  # 段后间距
    # 设置字体样式
    for run in paragraph.runs:
        run.font.name = "Times New Roman"  # 设置字体为宋体
        run.font.size = Pt(12)  # 小四字号（12磅）
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 确保设置为宋体

    return paragraph


def add_pic_header(doc, text):
    # 插入段落
    paragraph = doc.add_paragraph(text)

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.space_before = Pt(0)  # 段前间距
    paragraph_format.space_after = Pt(6)  # 段后间距
    # 设置字体样式
    for run in paragraph.runs:
        run.font.name = "Times New Roman"  # 设置字体为宋体
        run.font.size = Pt(12)  # 小四字号（12磅）
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 确保设置为宋体

    return paragraph


def insert_image(doc, image_path, width=None, height=None):
    """
    插入图片，并支持设置宽度和高度，图片默认居中显示。

    :param doc: Document 对象
    :param image_path: 图片路径
    :param width: 图片宽度（可选，单位为英寸，默认不设置）
    :param height: 图片高度（可选，单位为英寸，默认不设置）
    """
    if not Path(image_path).exists():
        print(f"图片路径无效：{image_path}")
        return
    # 创建段落
    paragraph = doc.add_paragraph()

    # 在段落中插入图片
    run = paragraph.add_run()

    # 设置图片宽度和高度
    if width and height:
        run.add_picture(
            image_path, width=Inches(width), height=Inches(height)
        )  # 设置宽高
    elif width:
        run.add_picture(image_path, width=Inches(width))  # 只设置宽度
    elif height:
        run.add_picture(image_path, height=Inches(height))  # 只设置高度
    else:
        run.add_picture(image_path)  # 默认不设置，按原图大小插入

    # 设置图片所在段落居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中显示图片


def apply_heading_style_to_all_titles(doc):
    """
    遍历文档中的所有一级标题，应用标题样式。

    :param doc: Document 对象
    """
    for paragraph in doc.paragraphs:
        # 判断该段落是否为一级标题
        if paragraph.style.name == "Heading 1":  # 默认为一级标题的样式名为 "Heading 1"
            set_heading_style(paragraph, level=1)  # 设置标题样式
        elif paragraph.style.name == "Heading 2":
            set_heading_style(paragraph, level=2)
        # elif paragraph.style.name == "Normal":
        #     set_paragraph_style(paragraph)


def setup_styles(doc):
    # =====================================
    # 正文样式 (Normal)
    # =====================================
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"  # 英文字体
    normal_style.font.size = Pt(11)  # 字号
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    # 段落格式
    normal_style.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进
    normal_style.paragraph_format.line_spacing = 1.5  # 1.15倍行距
    normal_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐

    # =====================================
    # List样式 (List Paragraph)
    list_style = doc.styles["List"]
    list_style.font.name = "Times New Roman"  # 英文字体
    list_style.font.size = Pt(11)  # 字号
    list_style.font.color.rgb = RGBColor(0, 0, 0)
    # 段落格式
    list_style.paragraph_format.first_line_indent = Cm(0)  # 首行缩进
    list_style.paragraph_format.line_spacing = 1.5  # 1.15倍行距
    list_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐

    # =====================================
    # 标题样式 (Heading 1-3)
    # =====================================
    heading_configs = {
        "Heading 1": {
            "font": "Arial Black",
            "size": Pt(16),
            "bold": True,
            "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "space_before": Pt(24),
            "space_after": Pt(12),
        },
        "Heading 2": {
            "font": "Arial Black",
            "size": Pt(14),
            "bold": True,
            "italic": True,
            "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "space_before": Pt(18),
            "space_after": Pt(6),
        },
        "Heading 3": {
            "font": "Times New Roman",
            "size": Pt(12),
            "bold": False,
            "italic": True,
            "underline": False,
            "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "space_before": Pt(12),
            "space_after": Pt(3),
        },
    }

    for style_name, config in heading_configs.items():
        style = doc.styles[style_name]
        # 字体设置
        style.font.name = config["font"]
        style.font.size = config["size"]
        style.font.bold = config.get("bold", False)
        style.font.italic = config.get("italic", False)
        style.font.underline = config.get("underline", False)
        style.font.color.rgb = RGBColor(0, 0, 0)  # 统一标题颜色
        # 段落格式
        style.paragraph_format.alignment = config["alignment"]
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = config["space_before"]
        style.paragraph_format.space_after = config["space_after"]
        style.paragraph_format.first_line_indent = Pt(0)  # 首行缩进 0 字符


def auto_report_EN():
    # 创建文档
    doc = Document()
    setup_styles(doc)  # 设置默认样式

    # 封面标题

    # 添加空段落，使标题位于页面中部
    for _ in range(6):  # 根据页面长度适当调整数量
        i = doc.add_paragraph()
        i.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = doc.add_paragraph()
    title_run = title.add_run(
        "XXX Petroleum Hydrocarbon Contaminated Site MIM Investigation Report"
    )
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    title_run.font.name = "Times New Roman"  # 设置中文字
    title_run.font.size = Pt(20)  # 设置字体大小为三号
    title_run.bold = True  # 加粗

    # 添加空段落，使标题位于页面中部
    for _ in range(10):  # 根据页面长度适当调整数量
        i = doc.add_paragraph()
        i.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    organization_paragraph = doc.add_paragraph()
    organization_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    organization_run = organization_paragraph.add_run(
        "XX environmental technology company"
    )
    organization_run.font.name = "Times New Roman"
    organization_run.font.size = Pt(12)
    r = organization_run._element  # 处理中文字体
    r.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # 添加时间
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    date_run = date_paragraph.add_run(datetime.now().strftime("%Y-%m"))
    date_run.font.name = "Times New Roman"
    date_run.font.size = Pt(16)
    # 添加分节符
    doc.add_section(WD_SECTION.NEW_PAGE)

    # 添加正文示例
    doc.add_heading("0 Clarification", level=1)
    doc.add_paragraph(
        "Bound by a confidentiality agreement, this report only provides examples and methods of arrangement, and does not contain real content.",
    )
    doc.add_paragraph(
        "Below we show how to insert multi-level headings, paragraphs, tables, images, etc. through a partial demo. Adjust the code according to different needs to generate a report that meets the requirements.",
    )
    doc.add_heading("1 Introduction", level=1)
    doc.add_paragraph(
        "A complete investigation report typically includes the following components:",
        style="List",
    )
    doc.add_paragraph(
        "1. Project Background: Brief introduction to the background and objectives of the investigation project",
        style="List",
    )
    doc.add_paragraph(
        "2. Field Investigation: Description of the basic site conditions, investigation methods, and investigation results",
        style="List",
    )
    doc.add_paragraph(
        "3. Data Analysis: Evaluation and analysis of investigation data to identify pollution source areas and contaminant plumes",
        style="List",
    )
    doc.add_paragraph(
        "4. Result Presentation: Display of investigation results through diagrams, tables, and other visual formats",
        style="List",
    )
    doc.add_paragraph(
        "5. Conclusion and Recommendations: Summary of investigation findings and proposals for follow-up actions",
        style="List",
    )

    doc.add_heading("2 Adding different levels of headings", level=1)
    doc.add_heading("2.1 Secondary headings", level=2)
    doc.add_heading("2.1.1 Tertiary heading", level=3)

    # 填充表格
    import random

    table_header = [
        "ID",
        "Rn(Bq/m³)",
        "VOCs(ppb)",
        "CO2(ppm)",
        "O2(%)",
        "CH4(%)",
        "FG(copies/g)",
    ]
    add_table_header(doc, "Table4-1 Microperturbation survey site data")
    table = doc.add_table(rows=7, cols=15)
    table.style = "Table Grid"
    for i, col_name in enumerate(table_header):
        table.cell(0, i).text = col_name
        for j in range(1, 15):
            table.cell(0, j).text = str(random.randint(10, 100))

    all_image_paths = [
        "./auto_report_cache/KMEANS-Radon.png",
        "./auto_report_cache/KMEANS-VOCs.png",
        "./auto_report_cache/KMEANS-CH4.png",
        "./auto_report_cache/KMEANS-CO2.png",
        "./auto_report_cache/KMEANS-O2.png",
        "./auto_report_cache/KMEANS-FG.png",
    ]
    image_paths = []
    for path in all_image_paths:
        if Path(path).exists():
            image_paths.append(path)
    # 计算表格的行数
    print(image_paths)
    rows = (len(image_paths) + 2 - 1) // 2

    # 创建表格
    table = doc.add_table(rows=rows, cols=2)

    # 填充表格单元格
    for i, image_path in enumerate(image_paths):
        row = i // 2
        col = i % 2
        cell = table.cell(row, col)
        if Path(image_path).exists():
            cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(2))
        else:
            cell.add_paragraph()
    labels = ["氡气", "VOCs", "CO2", "O2", "CH4", "功能基因"]
    all_image_paths = [
        "./auto_report_cache/Rn.png",
        "./auto_report_cache/VOCs.png",
        "./auto_report_cache/CH4.png",
        "./auto_report_cache/CO2.png",
        "./auto_report_cache/O2.png",
        "./auto_report_cache/FG.png",
    ]
    image_paths = []
    for path in all_image_paths:
        if Path(path).exists():
            image_paths.append(path)
    # 计算表格的行数
    print(image_paths)
    rows = (len(image_paths) + 2 - 1) // 2

    # 创建表格
    table = doc.add_table(rows=rows, cols=2)

    # 填充表格单元格
    for i, image_path in enumerate(image_paths):
        row = i // 2
        col = i % 2
        cell = table.cell(row, col)
        if Path(image_path).exists():
            cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(2))
        else:
            cell.add_paragraph()
    doc.add_paragraph("")

    # 应用样式
    return doc


if __name__ == "__main__":
    file = auto_report_EN()
    file.save("C:\\Users\\apple\\Desktop\\auto_report.docx")
