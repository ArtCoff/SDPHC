from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from pathlib import Path
from PIL import Image
from datetime import datetime

#
import os
import sys
import pandas as pd
import geopandas as gpd

# 获取当前脚本所在目录的上一级目录（即项目根目录）
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.append(project_root)
from utils.auto_report_EN import (
    add_table,
    add_bullet_list,
    insert_image,
    setup_styles,
    save_docx_safely,
)

print(f"Project root directory: {project_root}")
print(f"Current script directory: {Path.cwd()}")
cache_dir = Path(project_root) / "cache"
if not cache_dir.exists():
    cache_dir.mkdir(parents=True)


# 添加封面
def add_cover_page(doc):
    for _ in range(6):  # 根据页面长度适当调整数量
        i = doc.add_paragraph()
        i.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = doc.add_heading(
        "Non-Invasive Survey Report for Petroleum Hydrocarbon (PHC) Contaminated Site",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "Case Study: LNAPL Characterization at XXX Industrial Site"
    )
    subtitle.style = "Subtitle"
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):  # 根据页面长度适当调整数量
        i = doc.add_paragraph()
        i.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Prepared by: XYZ Environmental Consulting Firm").alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT
    )
    now = datetime.now()
    date_str = f"{now.strftime('%B')} {now.day}, {now.year}"  # 去除日期前导零
    doc.add_paragraph(f"Date: {date_str}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()


def add_Disclaimer(doc):
    doc.add_heading("Disclaimer", level=1)
    doc.add_paragraph(
        """
    This report is a simulated example created solely for academic illustration in the context of environmental site assessments. All data, including site descriptions, contamination levels, and regulatory references, are fictional or randomized to protect confidentiality and should not be interpreted as real-world measurements. The content follows general principles from international standards such as ASTM E1527 (Phase I Environmental Site Assessment) and EPA OSWER Directive, but no legal or technical verification of its accuracy has been performed. This document is intended exclusively for educational purposes and does not constitute a substitute for professional environmental investigations or compliance with actual regulatory requirements.
    
        """
    )
    doc.add_page_break()


# 添加执行摘要
def add_executive_summary(doc):
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "A Non-Invasive Survey (NIS) was conducted at the XXX industrial site to characterize LNAPL contamination from historical petroleum storage tank leaks. Key findings include:"
    )
    objectives = [
        "Contaminant Source Zone: Identified in the northern quadrant (Area A) via elevated VOCs (benzene up to 800 ppb) and radon-deficit zones.",
        "Pollution Extent: TPH (C6-C40) contamination spans 2.5 ha, with groundwater plume migration southwestward (Fig. 3).",
        "Biogeochemical Indicators: High *alkB* gene abundance (up to 1.2×10⁶ copies/g soil) confirms active hydrocarbon degradation.",
        "Recommendations: Prioritize Area A for detailed intrusive investigation and consider bioremediation enhanced by electron donor amendments.",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Number")
    doc.add_page_break()


# 添加引言章节（第2节）
def add_introduction_section(doc):
    doc.add_heading("2. Introduction", level=1)

    # 背景
    doc.add_heading("2.1 Background", level=2)
    doc.add_paragraph(
        "The XXX industrial site, located at 40°N, 75°W, operated as a petroleum storage facility from 2003. Historical records indicate multiple storage tank leaks, leading to potential LNAPL contamination of the shallow aquifer. This survey aims to characterize the contamination extent and assess associated risks."
    )

    # 调查目标
    doc.add_heading("2.2 Objectives", level=2)
    objectives = [
        "Map LNAPL source zones and subsurface distribution.",
        "Quantify VOC fluxes and biogeochemical degradation indicators.",
        "Assess risks to groundwater receptors and nearby residential areas.",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Bullet")

    doc.add_page_break()


# 添加法规框架章节（第3节）
def add_regulatory_framework_section(doc):
    doc.add_heading("3. Regulatory Framework", level=1)

    # 适用法规
    doc.add_heading("3.1 Applicable Standards", level=2)
    standards = [
        "EPA Risk Screening Levels (RSLs, 2023)**: Benzene (0.1 ppb), TPH (500 mg/kg soil, 100 µg/L groundwater).",
        "State-specific Vapor Intrusion Thresholds**: California DTSC (10 ppb benzene in soil gas).",
        "Risk Models: EPA RBCA (Tier 1) and EU Soil Screening Values (SSVs).",
    ]
    for std in standards:
        doc.add_paragraph(std, style="List Bullet")

    # 风险模型说明
    doc.add_heading("3.2 Risk Assessment Models", level=2)
    doc.add_paragraph(
        "Risk assessment follows EPA RBCA Tier 1 methodology, integrating exposure pathways (soil gas inhalation, groundwater ingestion) and receptor sensitivity. EU SSVs are used for cross-validation."
    )

    doc.add_page_break()


# 添加场地描述章节（第4节）
def add_site_description_section(doc):
    doc.add_heading("4. Site Description", level=1)

    # 物理特征
    doc.add_heading("4.1 Physical Characteristics", level=2)
    doc.add_paragraph(
        "The XXX site covers an area of 12 hectares (120,000 m²) and features sandy loam soil. The terrain is flat, with an elevation ranging from 10 to 15 meters above sea level. The site is actively engaged in industrial production, primarily manufacturing phenol and acetone , with ongoing operational activities at the facility."
    )

    # 水文地质
    doc.add_heading("4.2 Hydrogeology", level=2)
    hydrogeology = [
        "Shallow unconfined aquifer at depths of 3–8 meters.",
        "Groundwater flow direction: Southwest (SW), influenced by regional hydraulic gradient.",
        "Hydraulic conductivity: 10⁻⁴ cm/s (estimated from soil texture).",
    ]
    for item in hydrogeology:
        doc.add_paragraph(item, style="List Bullet")

    # 土地利用与受体
    doc.add_heading("4.3 Land Use and Receptors", level=2)
    doc.add_paragraph(
        "The site is located within an urban industrial park, with multiple neighboring factories operating in close proximity (approximately 500 meters to the east). The specific production activities of these facilities are currently unknown. A creek flows 200 meters to the south of the site, which may act as a potential surface water receptor for contaminant migration."
    )

    doc.add_page_break()


# 添加方法章节（含子章节）
def add_methodology_section(doc):
    doc.add_heading("5. Methodology", level=1)

    # 子章节
    doc.add_heading("5.1 Non-Invasive Survey (NIS) Approach", level=2)
    doc.add_paragraph(
        "The NIS approach integrates multi-parameter monitoring and biogeochemical mechanism analysis, as illustrated in Fig. 1."
    )
    insert_image(
        doc=doc,
        image_path=Path(project_root) / "assets/figs/conceptFig.png",
        width=12,
        figure_title="Figure 1: NIS Conceptual Framework",
    )
    # 子章节 5.1.1
    doc.add_heading("5.1.1 VOC Flux Monitoring", level=3)
    p = doc.add_paragraph()
    p.add_run("Instrument: ").bold = True
    p.add_run(
        "PID (Photoionization Detector, Ion Science Tiger Select) for BTEX and HCHO."
    )
    p = doc.add_paragraph()
    p.add_run("Protocol: ").bold = True
    p.add_run(
        "Grid spacing 20 m × 20 m; soil gas sampling depth 0.5 m using stainless steel probes."
    )
    doc.add_heading("5.1.2 CO₂/O₂ Flux Monitoring", level=3)
    # CO2/O2 测量
    co2o2_paragraph = doc.add_paragraph()
    co2o2_paragraph.add_run("Instrument: ").bold = True
    co2o2_paragraph.add_run(
        "Li-Cor 8100 soil gas analyzer with an infrared gas analyzer (IRGA) for real-time CO₂ and O₂ concentration measurements."
    )
    co2o2_paragraph = doc.add_paragraph()
    co2o2_paragraph.add_run("Protocol: ").bold = True
    content = [
        "Soil gas sampling depth: 0.5 m using stainless steel probes.",
        "Closed-chamber method: A 1 L chamber is placed over the probe, and CO₂/O₂ flux is calculated based on concentration gradient over time.",
        "Frequency: Bi-weekly measurements during survey period.",
    ]
    add_bullet_list(doc, items=content)
    # H2S/CH4 测量
    doc.add_heading("5.1.3 H₂S and CH₄ Detection", level=3)
    h2sch4_paragraph = doc.add_paragraph()
    h2sch4_paragraph.add_run("Instrument: ").bold = True
    content = [
        "H₂S: Jerome J605 portable hydrogen sulfide analyzer (electrochemical sensor, detection range: 0.1–200 ppm).",
        "CH₄: Los Gatos Research Ultra-Portable Greenhouse Gas Analyzer (laser absorption spectroscopy, detection limit: 1 ppb).",
    ]
    add_bullet_list(doc, items=content)
    h2sch4_paragraph = doc.add_paragraph()
    h2sch4_paragraph.add_run("Protocol: ").bold = True
    content = [
        "Soil gas sampling depth: 0.5 m using stainless steel probes.",
        "Real-time data logging with field calibration using certified gas standards (10 ppm H₂S, 5 ppm CH₄).",
    ]
    add_bullet_list(doc, items=content)
    doc.add_heading("5.1.4 Biogeochemical Indicators", level=3)
    p = doc.add_paragraph()
    p.add_run("Gas Analysis: ").bold = True
    doc.add_paragraph(
        "O₂/CO₂ flux via Li-Cor 8100 soil gas probes; H₂ and CH₄ via GC-TCD."
    )
    p = doc.add_paragraph()
    p.add_run("Microbial functional Gene Assays: ").bold = True
    doc.add_paragraph(
        " qPCR targeting C12O (aromatic hydrocarbon degradation), alkB (alkane oxidation), and prmA (methanogenesis/methylamine utilization)."
    )

    # 子章节 5.1.3
    doc.add_heading("5.1.5 Radon Deficit Mapping", level=3)
    p = doc.add_paragraph()
    p.add_run("Technique: ").bold = True
    p.add_run("AlphaGuard radon detector for near-surface ²²²Rn activity (Bq/m³).")
    doc.add_page_break()


# 添加结果章节（含表格和图表）
def add_results_section(doc, gdf):
    doc.add_heading("6. Results and Data Analysis", level=1)
    doc.add_heading("6.1 NIS survey", level=2)
    doc.add_paragraph(
        "The northern portion of the park contains the main production facilities (e.g., distillation and oxygenation workshops), while the southern portion includes storage, electrical, and fire protection facilities. Groundwater monitoring data indicate that PHCs may be leaking from the site. Taking into account the actual production and facility distribution in the park, a total of 146 NIS monitoring points were deployed, including 29 functional gene measurement points and 15 radon measurement points (see Figure 2 for specific locations). Based on the NIS monitoring data, Empirical threshold analysis was used to carry out the investigation of the current status of PHCs pollution."
    )

    insert_image(
        doc,
        image_path=Path(project_root) / "assets/figs/NIS_point.jpg",
        width=10,
        figure_title="Figure 2: Sampling Point Distribution",
        legend="Legend: Dots = NIS survey point; Black closed polygon = Park Boundary",
    )
    nis_data = gdf[["point_code", "VOCs", "O2", "CO2", "CH4", "H2S"]].copy()
    radon_data = gdf[["point_code", "Radon"]].copy().dropna()
    exceed_data = gdf[gdf["The_other_soil_gas_scores"] >= 6].copy()
    exceed_data = exceed_data[
        [
            "Point_ID",
            "VOCs_Score",
            "CO2_Score",
            "H2_Score",
            "O2_Score",
            "CH4_Score",
            "H2S_Score",
            "The_other_soil_gas_scores",
        ]
    ]
    doc.add_paragraph(
        "Table 1 presents the NIS sampling data from the chemical industrial park, comprising 149 sampling points. The monitored parameters include volatile organic compounds (VOCs), oxygen (O₂), carbon dioxide (CO₂), methane (CH₄), and hydrogen sulfide (H₂S), among others. The radon gas sampling data are shown in Table 2."
    )
    add_table(doc, df=nis_data, table_title="Table 1: NIS survey points data")
    add_table(
        doc,
        df=radon_data,
        table_title="Table 2: Radon gas sampling data",
    )
    doc.add_heading("6.2 Analysis results", level=2)
    doc.add_paragraph(
        "An empirical threshold analysis was conducted to identify sampling points exceeding regulatory benchmarks (Fig. 1). Using EPA Risk Screening Levels (RSLs) for benzene (0.1 ppb) and DTSC thresholds (10 ppb for soil gas), 12 out of 45 points (26.7%) were classified as high-risk zones. Notably, Sample S-08 in Area A exhibited benzene concentrations 850 times higher than the EPA RSL. Microbial gene assays further confirmed contamination: the C12O gene (aromatic hydrocarbon degradation) showed abundance peaks of 9.6×10⁵ copies/g soil in these zones, directly linking VOC anomalies to petroleum hydrocarbon degradation pathways."
    )
    add_table(doc, df=exceed_data, table_title="Table 3: Exceedance points")
    doc.add_paragraph()
    doc.add_paragraph(
        "Fig.3 shows the manual delineation based on the anomalies of each monitoring point to determine the pollution situation, and finally screened out the pollution source area and the suspected pollution source area, according to the delineation of the scope to determine the main leakage area is located in the Loading platform and Hazardous waste room (red circled area)), and the orange circled area is also assigned to the key monitoring area. According to the delineation of the scope, the main leakage areas are located in the Loading platform and Hazardous waste room (circled in red), and the area circled in orange is also included in the key monitoring area. Fig.4 shows the contamination range, which was evaluated based on the anomaly scores of all the monitoring points ((anomalies are marked in red)), and the final result shows that the whole plant is contaminated, and the contamination may escape from the boundary, except for the southwest boundary."
    )
    insert_image(
        doc,
        image_path=Path(project_root) / "assets/figs/sourcezone.jpg",
        width=10,
        figure_title="Figure 3: Source Zone Identification",
        legend="Legend: Red = LNAPL source zone; Orange = Suspected LNAPL zone",
    )
    insert_image(
        doc,
        image_path=Path(project_root) / "assets/figs/scope.png",
        width=10,
        figure_title="Figure 4: Contamination Scope",
        legend="Legend: Red = Contamination Scope; White = Clean Soil",
    )
    doc.add_page_break()


def add_risk_assessment_section(doc):
    doc.add_heading("7. Risk Assessment", level=1)

    # 7.1 Human Health Risks
    doc.add_heading("7.1 Human Health Risks", level=2)
    doc.add_paragraph(
        "The risk assessment follows the EPA RBCA Tier 1 methodology, integrating exposure pathways and receptor sensitivity."
    )

    # 子章节：暴露途径
    doc.add_heading("7.1.1 Exposure Pathways", level=3)
    exposure_pathways = [
        "Inhalation of soil gas (vapor intrusion into residential buildings).",
        "Dermal contact with contaminated soil.",
        "Ingestion of contaminated groundwater.",
    ]
    for path in exposure_pathways:
        doc.add_paragraph(path, style="List Bullet")

    # 子章节：受体分析
    doc.add_heading("7.1.2 Receptor Analysis", level=3)
    doc.add_paragraph(
        "The primary receptors include nearby residents living within 500 meters east of the site and potential future industrial workers on-site."
    )

    # 子章节：风险计算
    doc.add_heading("7.1.3 Risk Calculation", level=3)
    risk_table_data = [
        [
            "Contaminant",
            "Carcinogenic Risk",
            "Non-Carcinogenic Risk (HQ)",
            "Exceeds Threshold?",
        ],
        ["Benzene", "1.2×10⁻⁴", "0.8", "Yes"],
        ["TPH (C6-C40)", "N/A", "1.5", "Yes"],
        ["Toluene", "N/A", "0.6", "No"],
    ]
    add_table(
        doc,
        df=pd.DataFrame(risk_table_data, columns=risk_table_data[0]),
        table_title="Table 2: Risk Assessment Summary",
    )
    doc.add_paragraph("Threshold: Carcinogenic Risk >1×10⁻⁶ or HQ >1.0.")

    # 7.2 Ecological Risks
    doc.add_heading("7.2 Ecological Risks", level=2)
    doc.add_paragraph(
        "Ecological risks were assessed using the EU Soil Screening Values (SSVs) and EPA ECO-SSL guidelines."
    )

    # 子章节：土壤生态影响
    doc.add_heading("7.2.1 Soil Toxicity", level=3)
    soil_toxicity = [
        "TPH >10,000 mg/kg in Area A inhibits microbial diversity.",
        "Reduced earthworm survival rate observed in laboratory toxicity tests.",
    ]
    for item in soil_toxicity:
        doc.add_paragraph(item, style="List Bullet")

    # 子章节：地表水生态影响
    doc.add_heading("7.2.2 Surface Water Impact", level=3)
    doc.add_paragraph(
        "The plume discharge into the nearby creek may affect benthic organisms. Predicted TPH concentrations exceed EPA aquatic toxicity thresholds."
    )
    # 7.3 Risk Summary
    doc.add_heading("7.3 Risk Summary", level=2)
    doc.add_paragraph("The risk assessment concludes that:")
    risk_summary = [
        "Benzene and TPH pose unacceptable carcinogenic and non-carcinogenic risks to nearby residents.",
        "Groundwater contamination threatens drinking water supplies and aquatic ecosystems.",
        "Immediate mitigation measures are required, including vapor intrusion controls and source zone remediation.",
    ]
    for item in risk_summary:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()


# 添加结论章节
def add_conclusion_section(doc):
    doc.add_heading("8. Conclusions and Recommendations", level=1)
    doc.add_paragraph(
        "The NIS successfully identified the LNAPL source zone (Area A) with active volatilization and biodegradation processes."
    )
    doc.add_paragraph("Recommendations:")
    conclusions = [
        "Phase III Investigation: Install monitoring wells in Area A for LNAPL thickness quantification."
        "Remediation Options: Bioremediation with nitrate amendment to enhance aromatic degradation."
        "Long-Term Monitoring: Bi-monthly VOC and microbial functional gene abundance tracking."
    ]
    for item in conclusions:
        doc.add_paragraph(item, style="List Number")


# 主函数
def auto_report_for_empirical_threshold_analysis(gdf=None):
    doc = Document()
    setup_styles(doc)  # 设置默认样式

    add_cover_page(doc)
    add_Disclaimer(doc)
    # add_table_of_contents(doc)
    add_executive_summary(doc)
    add_introduction_section(doc)
    add_regulatory_framework_section(doc)
    add_site_description_section(doc)
    add_methodology_section(doc)
    add_results_section(doc, gdf)
    add_risk_assessment_section(doc)
    add_conclusion_section(doc)

    return doc


if __name__ == "__main__":
    import geopandas as gpd

    gdf = gpd.read_file("C:\\Users\\Apple\\Desktop\\SDPHC\\tests\\LX_ETA.gpkg")
    print(gdf.head())
    print(gdf.columns)
    file = auto_report_for_empirical_threshold_analysis(gdf=gdf)
    save_docx_safely(file, "C:\\Users\\Apple\\Desktop\\AppendixB.docx")
