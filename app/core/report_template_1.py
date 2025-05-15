from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from pathlib import Path
from PIL import Image
from datetime import datetime
from utils.auto_report_EN import (
    set_heading_style,
    set_paragraph_style,
    add_pic_header,
    add_table_header,
    insert_image,
    setup_styles,
    save_docx_safely,
)


# 添加封面
def add_cover_page(doc):
    for _ in range(6):  # 根据页面长度适当调整数量
        i = doc.add_paragraph()
        i.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = doc.add_heading(
        "Non-Invasive Survey Report for Petroleum Hydrocarbon (PHC) Contaminated Site",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "Case Study: LNAPL Characterization at XXX Industrial Site"
    )
    subtitle.style = "Subtitle"
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):  # 根据页面长度适当调整数量
        i = doc.add_paragraph()
        i.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Prepared by: XYZ Environmental Consulting Firm").alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT
    )
    now = datetime.now()
    date_str = f"{now.strftime('%B')} {now.day}, {now.year}"  # 去除日期前导零
    doc.add_paragraph(f"Date: {date_str}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()


# 添加执行摘要
def add_executive_summary(doc):
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "A Non-Invasive Survey (NIS) was conducted at the XXX industrial site to characterize LNAPL contamination from historical petroleum storage tank leaks. Key findings include:"
    )
    objectives = [
        "Contaminant Source Zone: Identified in the northern quadrant (Area A) via elevated VOCs (benzene up to 850 ppb) and radon-deficit zones.",
        "Pollution Extent: TPH (C6-C40) contamination spans 2.5 ha, with groundwater plume migration southwestward (Fig. 3).",
        "Biogeochemical Indicators: High *alkB* gene abundance (up to 1.2×10⁶ copies/g soil) confirms active hydrocarbon degradation.",
        "Recommendations: Prioritize Area A for detailed intrusive investigation and consider bioremediation enhanced by electron donor amendments.",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Number")
    doc.add_page_break()


# 添加引言章节（第2节）
def add_introduction_section(doc):
    doc.add_heading("2. Introduction", level=1)

    # 背景
    doc.add_heading("2.1 Background", level=2)
    doc.add_paragraph(
        "The XXX industrial site, located at 40°N, 75°W, operated as a petroleum storage facility from 1960 to 1998. Historical records indicate multiple underground storage tank (UST) leaks, leading to potential LNAPL contamination of the shallow aquifer. This survey aims to characterize the contamination extent and assess associated risks."
    )

    # 调查目标
    doc.add_heading("2.2 Objectives", level=2)
    objectives = [
        "Map LNAPL source zones and subsurface distribution.",
        "Quantify VOC fluxes and biogeochemical degradation indicators.",
        "Assess risks to groundwater receptors and nearby residential areas.",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Bullet")

    doc.add_page_break()


# 添加法规框架章节（第3节）
def add_regulatory_framework_section(doc):
    doc.add_heading("3. Regulatory Framework", level=1)

    # 适用法规
    doc.add_heading("3.1 Applicable Standards", level=2)
    standards = [
        "EPA Risk Screening Levels (RSLs, 2023)**: Benzene (0.1 ppb), TPH (500 mg/kg soil, 100 µg/L groundwater).",
        "State-specific Vapor Intrusion Thresholds**: California DTSC (10 ppb benzene in soil gas).",
        "Risk Models**: EPA RBCA (Tier 1) and EU Soil Screening Values (SSVs).",
    ]
    for std in standards:
        doc.add_paragraph(std, style="List Bullet")

    # 风险模型说明
    doc.add_heading("3.2 Risk Assessment Models", level=2)
    doc.add_paragraph(
        "Risk assessment follows EPA RBCA Tier 1 methodology, integrating exposure pathways (soil gas inhalation, groundwater ingestion) and receptor sensitivity. EU SSVs are used for cross-validation."
    )

    doc.add_page_break()


# 添加场地描述章节（第4节）
def add_site_description_section(doc):
    doc.add_heading("4. Site Description", level=1)

    # 物理特征
    doc.add_heading("4.1 Physical Characteristics", level=2)
    doc.add_paragraph(
        "The XXX site spans 15 hectares with sandy loam soil texture. The terrain is flat with an elevation range of 10–15 m above sea level. No active operations are currently present on-site."
    )

    # 水文地质
    doc.add_heading("4.2 Hydrogeology", level=2)
    hydrogeology = [
        "Shallow unconfined aquifer at depths of 3–8 meters.",
        "Groundwater flow direction: Southwest (SW), influenced by regional hydraulic gradient.",
        "Hydraulic conductivity: 10⁻⁴ cm/s (estimated from soil texture).",
    ]
    for item in hydrogeology:
        doc.add_paragraph(item, style="List Bullet")

    # 土地利用与受体
    doc.add_heading("4.3 Land Use and Receptors", level=2)
    doc.add_paragraph(
        "The site is bordered by a residential zone 500 meters to the east. A creek flows 200 meters to the south, serving as a potential surface water receptor for contaminant migration."
    )

    doc.add_page_break()


# 添加方法章节（含子章节）
def add_methodology_section(doc):
    doc.add_heading("5. Methodology", level=1)

    # 子章节
    doc.add_heading("5.1 Non-Invasive Survey (NIS) Approach", level=2)
    doc.add_paragraph(
        "The NIS approach integrates multi-parameter monitoring and biogeochemical mechanism analysis, as illustrated in Fig. 1."
    )

    # 子章节 5.1.1
    doc.add_heading("5.1.1 VOC Flux Monitoring", level=3)
    p = doc.add_paragraph()
    p.add_run("Instrument: ").bold = True
    p.add_run(
        "PID (Photoionization Detector, Ion Science Tiger Select) for BTEX and HCHO."
    )
    p = doc.add_paragraph()
    p.add_run("Protocol: ").bold = True
    p.add_run(
        "Grid spacing 20 m × 20 m; soil gas sampling depth 0.5 m using stainless steel probes."
    )

    # 子章节 5.1.2
    doc.add_heading("5.1.2 Biogeochemical Indicators", level=3)
    p = doc.add_paragraph()
    p.add_run("Gas Analysis: ").bold = True
    p.add_run("O₂/CO₂ flux via Li-Cor 8100 soil gas probes; H₂ and CH₄ via GC-TCD.")
    p = doc.add_paragraph()
    p.add_run("Microbial Gene Assays: ").bold = True
    p.add_run(
        "qPCR targeting *alkB* (alkanes), *todC1* (aromatics), and *mcrA* (methanogens)."
    )

    # 子章节 5.1.3
    doc.add_heading("5.1.3 Radon Deficit Mapping", level=3)
    p = doc.add_paragraph()
    p.add_run("Technique: ").bold = True
    p.add_run("AlphaGuard radon detector for near-surface ²²²Rn activity (Bq/m³).")
    p = doc.add_paragraph()
    p.add_run("Interpretation: ").bold = True
    p.add_run(
        "Background 25–30 Bq/m³; deficit zones (<10 Bq/m³) indicate NAPL presence."
    )

    doc.add_page_break()


# 添加结果章节（含表格和图表）
def add_results_section(doc):
    doc.add_heading("6. Results and Data Analysis", level=1)

    # 图表占位符
    # doc.add_heading("Figure 1: Sampling Point Distribution", level=2)
    # doc.add_picture("sampling_map.png", width=Inches(5.5))  # 替换为实际图片路径
    # insert_image()
    add_pic_header(doc, "Figure 1: Sampling Point Distribution")
    caption = doc.add_paragraph(
        "Legend: Red dots = VOC anomalies; Blue triangles = radon deficit zones"
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表格
    # doc.add_heading("Table 1: VOC Concentrations (ppb)", level=2)
    add_table_header(doc, "Table 1: VOC Concentrations (ppb)")
    data = [
        ["Sample ID", "X (m)", "Y (m)", "Benzene", "Toluene", "HCHO", "PID Response"],
        ["S-01", "20", "20", "12", "8", "5", "150"],
        ["S-08", "100", "60", "850", "620", "30", "2100"],
        ["S-15", "180", "100", "45", "30", "10", "300"],
    ]
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = data[i][j]
    doc.add_page_break()


def add_risk_assessment_section(doc):
    doc.add_heading("7. Risk Assessment", level=1)

    # 7.1 Human Health Risks
    doc.add_heading("7.1 Human Health Risks", level=2)
    doc.add_paragraph(
        "The risk assessment follows the EPA RBCA Tier 1 methodology, integrating exposure pathways and receptor sensitivity."
    )

    # 子章节：暴露途径
    doc.add_heading("7.1.1 Exposure Pathways", level=3)
    exposure_pathways = [
        "Inhalation of soil gas (vapor intrusion into residential buildings).",
        "Dermal contact with contaminated soil.",
        "Ingestion of contaminated groundwater.",
    ]
    for path in exposure_pathways:
        doc.add_paragraph(path, style="List Bullet")

    # 子章节：受体分析
    doc.add_heading("7.1.2 Receptor Analysis", level=3)
    doc.add_paragraph(
        "The primary receptors include nearby residents living within 500 meters east of the site and potential future industrial workers on-site."
    )

    # 子章节：风险计算
    doc.add_heading("7.1.3 Risk Calculation", level=3)
    risk_table_data = [
        [
            "Contaminant",
            "Carcinogenic Risk",
            "Non-Carcinogenic Risk (HQ)",
            "Exceeds Threshold?",
        ],
        ["Benzene", "1.2×10⁻⁴", "0.8", "Yes"],
        ["TPH (C6-C40)", "N/A", "1.5", "Yes"],
        ["Toluene", "N/A", "0.6", "No"],
    ]
    risk_table = doc.add_table(rows=len(risk_table_data), cols=len(risk_table_data[0]))
    risk_table.style = "Table Grid"
    for i, row in enumerate(risk_table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = risk_table_data[i][j]
    doc.add_paragraph("Threshold: Carcinogenic Risk >1×10⁻⁶ or HQ >1.0.")

    # 7.2 Ecological Risks
    doc.add_heading("7.2 Ecological Risks", level=2)
    doc.add_paragraph(
        "Ecological risks were assessed using the EU Soil Screening Values (SSVs) and EPA ECO-SSL guidelines."
    )

    # 子章节：土壤生态影响
    doc.add_heading("7.2.1 Soil Toxicity", level=3)
    soil_toxicity = [
        "TPH >10,000 mg/kg in Area A inhibits microbial diversity.",
        "Reduced earthworm survival rate observed in laboratory toxicity tests.",
    ]
    for item in soil_toxicity:
        doc.add_paragraph(item, style="List Bullet")

    # 子章节：地表水生态影响
    doc.add_heading("7.2.2 Surface Water Impact", level=3)
    doc.add_paragraph(
        "The plume discharge into the nearby creek may affect benthic organisms. Predicted TPH concentrations exceed EPA aquatic toxicity thresholds."
    )

    # 图表占位符
    # doc.add_heading("Figure 2: Risk Assessment Model Output", level=2)
    # doc.add_picture("risk_model_output.png", width=Inches(5.5))  # 替换为实际图片路径
    # insert_image()
    add_pic_header(doc, "Figure 2: Risk Assessment Model Output")
    caption = doc.add_paragraph(
        "Legend: Red zones = high cancer risk; Yellow zones = moderate non-carcinogenic risk."
    )

    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 7.3 Risk Summary
    doc.add_heading("7.3 Risk Summary", level=2)
    doc.add_paragraph("The risk assessment concludes that:")
    risk_summary = [
        "Benzene and TPH pose unacceptable carcinogenic and non-carcinogenic risks to nearby residents.",
        "Groundwater contamination threatens drinking water supplies and aquatic ecosystems.",
        "Immediate mitigation measures are required, including vapor intrusion controls and source zone remediation.",
    ]
    for item in risk_summary:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()


# 添加结论章节
def add_conclusion_section(doc):
    doc.add_heading("8. Conclusions and Recommendations", level=1)
    doc.add_paragraph(
        "The NIS successfully identified the LNAPL source zone (Area A) with active volatilization and biodegradation processes."
    )
    doc.add_paragraph("Recommendations:")
    conclusions = [
        "Phase III Investigation: Install monitoring wells in Area A for LNAPL thickness quantification."
        "Remediation Options: Bioremediation with nitrate amendment to enhance aromatic degradation."
        "Long-Term Monitoring: Bi-monthly VOC and microbial gene abundance tracking."
    ]
    for item in conclusions:
        doc.add_paragraph(item, style="List Number")


# 主函数
def auto_report_for_empirical_threshold_analysis():
    doc = Document()
    setup_styles(doc)  # 设置默认样式
    add_cover_page(doc)
    # add_table_of_contents(doc)
    add_executive_summary(doc)
    add_introduction_section(doc)
    add_regulatory_framework_section(doc)
    add_site_description_section(doc)
    add_methodology_section(doc)
    add_results_section(doc)
    add_risk_assessment_section(doc)
    add_conclusion_section(doc)

    return doc


if __name__ == "__main__":
    file = auto_report_for_empirical_threshold_analysis()
    save_docx_safely(file, "C:\\Users\\Apple\\Desktop\\auto_report.docx")
